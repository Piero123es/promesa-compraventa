"""
Generador de "Promesa de Compraventa" a partir de plantilla.docx
-----------------------------------------------------------------
Requisitos (instalar una sola vez):
    pip install customtkinter docxtpl

IMPORTANTE:
Este script usa la librería `docxtpl`, no `python-docx`, porque tu
plantilla.docx tiene etiquetas tipo Jinja tal cual: {{comprador}},
{{dni}}, {% for cuota in cuotas %} ... {% endfor %}, etc.
python-docx NO reemplaza esas etiquetas automáticamente, por eso
"faltaban datos": el formulario original solo pedía 7 campos y la
plantilla necesita muchos más (fecha, terreno, precio, cuotas,
contacto adicional y ejecutivo comercial).

NUEVO: después de rellenar la plantilla, el script fuerza que TODO
el texto del documento generado quede en fuente Aptos, tamaño 11
(cuerpo, tablas, encabezados, pies de página, y el estilo "Normal"
para cualquier texto que no tenga formato explícito).

Coloca este archivo en la misma carpeta que "plantilla.docx"
(o cambia la variable RUTA_PLANTILLA más abajo).
"""

import os
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docxtpl import DocxTemplate
from docx.shared import Pt
from docx.oxml.ns import qn

# -----------------------------
# Configuración general
# -----------------------------
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

RUTA_PLANTILLA = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantilla.docx")

NOMBRE_FUENTE = "Aptos"
TAMANO_FUENTE = 11


# -----------------------------
# Utilidades para forzar fuente en TODO el documento
# -----------------------------
def _set_run_font(run, nombre_fuente, tamano_pt, mayusculas=True):
    run.font.name = nombre_fuente
    run.font.size = Pt(tamano_pt)
    run.font.all_caps = mayusculas
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for atributo in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(atributo), nombre_fuente)


def _procesar_parrafos(parrafos, nombre_fuente, tamano_pt, mayusculas=True):
    for p in parrafos:
        for run in p.runs:
            _set_run_font(run, nombre_fuente, tamano_pt, mayusculas)


def _procesar_tablas(tablas, nombre_fuente, tamano_pt, mayusculas=True):
    for tabla in tablas:
        for fila in tabla.rows:
            for celda in fila.cells:
                _procesar_parrafos(celda.paragraphs, nombre_fuente, tamano_pt, mayusculas)
                _procesar_tablas(celda.tables, nombre_fuente, tamano_pt, mayusculas)  # tablas anidadas


def aplicar_fuente_global(documento, nombre_fuente=NOMBRE_FUENTE, tamano_pt=TAMANO_FUENTE, mayusculas=True):
    """Fuerza que TODO el texto del documento use la fuente y tamaño indicados,
    y lo muestra en MAYÚSCULAS (formato 'todo en mayúsculas' de Word, no cambia
    el texto real, solo cómo se ve/imprime): cuerpo, tablas, encabezados y pies
    de página. También actualiza el estilo 'Normal' para que cualquier texto sin
    formato explícito quede igual."""
    _procesar_parrafos(documento.paragraphs, nombre_fuente, tamano_pt, mayusculas)
    _procesar_tablas(documento.tables, nombre_fuente, tamano_pt, mayusculas)

    for section in documento.sections:
        partes = (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        )
        for parte in partes:
            _procesar_parrafos(parte.paragraphs, nombre_fuente, tamano_pt, mayusculas)
            _procesar_tablas(parte.tables, nombre_fuente, tamano_pt, mayusculas)

    # Estilo "Normal": afecta cualquier texto nuevo o sin formato explícito
    estilo_normal = documento.styles["Normal"]
    estilo_normal.font.name = nombre_fuente
    estilo_normal.font.size = Pt(tamano_pt)
    estilo_normal.font.all_caps = mayusculas
    rPr = estilo_normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for atributo in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(atributo), nombre_fuente)


# -----------------------------
# Fila de cuota (widget reutilizable)
# -----------------------------
class FilaCuota(ctk.CTkFrame):
    def __init__(self, master, on_eliminar):
        super().__init__(master, fg_color="#f0f0f0", corner_radius=8)

        self.nombre = ctk.CTkEntry(self, placeholder_text="Ej: Inicial / Cuota 1", width=160)
        self.nombre.grid(row=0, column=0, padx=5, pady=8)

        self.monto = ctk.CTkEntry(self, placeholder_text="Monto S/", width=110)
        self.monto.grid(row=0, column=1, padx=5, pady=8)

        self.monto_letras = ctk.CTkEntry(self, placeholder_text="Monto en letras", width=220)
        self.monto_letras.grid(row=0, column=2, padx=5, pady=8)

        self.fecha = ctk.CTkEntry(self, placeholder_text="dd/mm/aaaa", width=110)
        self.fecha.grid(row=0, column=3, padx=5, pady=8)

        self.btn_eliminar = ctk.CTkButton(
            self, text="✕", width=30, fg_color="#d9534f", hover_color="#c9302c",
            command=lambda: on_eliminar(self)
        )
        self.btn_eliminar.grid(row=0, column=4, padx=5, pady=8)

    def obtener_datos(self):
        return {
            "nombre": self.nombre.get().strip(),
            "monto": self.monto.get().strip(),
            "monto_letras": self.monto_letras.get().strip(),
            "fecha": self.fecha.get().strip(),
        }


# -----------------------------
# Formulario principal
# -----------------------------
class Formulario(ctk.CTkToplevel):
    def __init__(self, master):
        super().__init__(master)
        self.title("Nueva Promesa de Compraventa")
        self.geometry("950x750")

        self.filas_cuotas = []

        contenedor = ctk.CTkScrollableFrame(self, width=900, height=700)
        contenedor.pack(fill="both", expand=True, padx=10, pady=10)

        ctk.CTkLabel(contenedor, text="PROMESA DE COMPRAVENTA",
                     font=("Arial", 24, "bold")).pack(pady=(5, 15))

        # ---------- Fecha del documento ----------
        self._titulo_seccion(contenedor, "Fecha del documento")
        fila = ctk.CTkFrame(contenedor, fg_color="transparent")
        fila.pack(pady=5)
        self.dia = self._campo_horizontal(fila, "Día", 0, ancho=80)
        self.mes = self._campo_horizontal(fila, "Mes", 1, ancho=140)

        # ---------- Datos del comprador ----------
        self._titulo_seccion(contenedor, "Datos del Promitente Comprador(a)")
        self.comprador = self._campo(contenedor, "Nombre Completo")
        self.dni = self._campo(contenedor, "DNI")
        self.celular = self._campo(contenedor, "Celular")
        self.nacionalidad = self._campo(contenedor, "Nacionalidad")
        self.estado_civil = self._campo(contenedor, "Estado Civil")
        self.ocupacion = self._campo(contenedor, "Ocupación")
        self.direccion = self._campo(contenedor, "Dirección")
        self.provincia = self._campo(contenedor, "Provincia")
        self.departamento = self._campo(contenedor, "Departamento")

        # ---------- Datos del terreno / precio ----------
        self._titulo_seccion(contenedor, "Datos del Terreno y Precio")
        self.area = self._campo(contenedor, "Área (M2)")
        self.lote = self._campo(contenedor, "N° de Lote")
        self.precio = self._campo(contenedor, "Precio total (S/)")
        self.precio_letras = self._campo(contenedor, "Precio total en letras")
        self.monto_adicional = self._campo(contenedor, "Monto para escritura pública (S/)")
        self.monto_adicional_letras = self._campo(contenedor, "Monto para escritura pública (en letras)")

        # ---------- Cuotas ----------
        self._titulo_seccion(contenedor, "Cuotas de pago")
        self.frame_cuotas = ctk.CTkFrame(contenedor, fg_color="transparent")
        self.frame_cuotas.pack(fill="x", pady=5)

        ctk.CTkButton(contenedor, text="+ Agregar cuota",
                      command=self.agregar_cuota).pack(pady=(0, 15))
        self.agregar_cuota()  # una fila inicial

        # ---------- Contacto adicional ----------
        self._titulo_seccion(contenedor, "Contacto Adicional")
        self.adicional = self._campo(contenedor, "Nombre del contacto adicional")
        self.parentesco = self._campo(contenedor, "Parentesco")
        self.celular_adicional = self._campo(contenedor, "Celular del contacto adicional")

        # ---------- Ejecutivo comercial ----------
        self._titulo_seccion(contenedor, "Ejecutivo Comercial")
        self.ejecutivo = self._campo(contenedor, "Nombre del ejecutivo")
        self.dni_ejecutivo = self._campo(contenedor, "DNI del ejecutivo")

        # ---------- Botón generar ----------
        ctk.CTkButton(contenedor, text="Generar Word", height=45,
                      font=("Arial", 15, "bold"),
                      command=self.generar_word).pack(pady=25)

    # --- helpers de UI ---
    def _titulo_seccion(self, master, texto):
        ctk.CTkLabel(master, text=texto, font=("Arial", 16, "bold"),
                     text_color="#1f538d").pack(pady=(20, 5), anchor="w", padx=10)

    def _campo(self, master, etiqueta, ancho=500):
        ctk.CTkLabel(master, text=etiqueta).pack(anchor="w", padx=10)
        entry = ctk.CTkEntry(master, width=ancho)
        entry.pack(pady=(0, 8), padx=10)
        return entry

    def _campo_horizontal(self, master, etiqueta, columna, ancho=120):
        sub = ctk.CTkFrame(master, fg_color="transparent")
        sub.grid(row=0, column=columna, padx=10)
        ctk.CTkLabel(sub, text=etiqueta).pack()
        entry = ctk.CTkEntry(sub, width=ancho)
        entry.pack()
        return entry

    # --- cuotas dinámicas ---
    def agregar_cuota(self):
        fila = FilaCuota(self.frame_cuotas, self.eliminar_cuota)
        fila.pack(pady=3, fill="x")
        self.filas_cuotas.append(fila)

    def eliminar_cuota(self, fila):
        if len(self.filas_cuotas) <= 1:
            messagebox.showwarning("Aviso", "Debe haber al menos una cuota.")
            return
        fila.destroy()
        self.filas_cuotas.remove(fila)

    # --- generación del documento ---
    def generar_word(self):
        # Validación básica de campos obligatorios
        obligatorios = {
            "Nombre del comprador": self.comprador.get(),
            "DNI": self.dni.get(),
            "Día": self.dia.get(),
            "Mes": self.mes.get(),
            "Precio": self.precio.get(),
        }
        faltantes = [nombre for nombre, valor in obligatorios.items() if not valor.strip()]
        if faltantes:
            messagebox.showerror("Faltan datos", "Completa: " + ", ".join(faltantes))
            return

        if not os.path.exists(RUTA_PLANTILLA):
            messagebox.showerror(
                "Plantilla no encontrada",
                f"No se encontró plantilla.docx en:\n{RUTA_PLANTILLA}\n\n"
                "Coloca el archivo plantilla.docx junto a este script."
            )
            return

        cuotas = [f.obtener_datos() for f in self.filas_cuotas]

        contexto = {
            "dia": self.dia.get().strip(),
            "mes": self.mes.get().strip(),
            "comprador": self.comprador.get().strip(),
            "dni": self.dni.get().strip(),
            "celular": self.celular.get().strip(),
            "nacionalidad": self.nacionalidad.get().strip(),
            "ocupacion": self.ocupacion.get().strip(),
            "estado_civil": self.estado_civil.get().strip(),
            "direccion": self.direccion.get().strip(),
            "provincia": self.provincia.get().strip(),
            "departamento": self.departamento.get().strip(),
            "area": self.area.get().strip(),
            "lote": self.lote.get().strip(),
            "precio": self.precio.get().strip(),
            "precio_letras": self.precio_letras.get().strip(),
            "monto_adicional": self.monto_adicional.get().strip(),
            "monto_adicional_letras": self.monto_adicional_letras.get().strip(),
            "cuotas": cuotas,
            "adicional": self.adicional.get().strip(),
            "parentesco": self.parentesco.get().strip(),
            "celular_adicional": self.celular_adicional.get().strip(),
            "ejecutivo": self.ejecutivo.get().strip(),
            "dni_ejecutivo": self.dni_ejecutivo.get().strip(),
        }

        # Pedir dónde guardar
        nombre_sugerido = f"Promesa_Compraventa_{contexto['comprador'].replace(' ', '_') or 'documento'}.docx"
        ruta_salida = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=nombre_sugerido,
            filetypes=[("Documento Word", "*.docx")],
            title="Guardar documento generado"
        )
        if not ruta_salida:
            return  # el usuario canceló

        try:
            doc = DocxTemplate(RUTA_PLANTILLA)
            doc.render(contexto)
            # Forzar Aptos 11 en TODO el documento generado
            aplicar_fuente_global(doc.docx, NOMBRE_FUENTE, TAMANO_FUENTE)
            doc.save(ruta_salida)
        except Exception as e:
            messagebox.showerror("Error al generar", f"Ocurrió un error:\n{e}")
            return

        messagebox.showinfo("Listo", f"Documento generado correctamente en:\n{ruta_salida}")


# -----------------------------
# Ventana principal
# -----------------------------
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Sistema de Promesas")
        self.geometry("900x600")

        ctk.CTkLabel(self, text="PROMESA DE COMPRAVENTA",
                     font=("Arial", 30, "bold")).pack(pady=50)

        ctk.CTkButton(self, text="Nuevo Documento", width=250, height=50,
                      command=self.nuevo_documento).pack()

    def nuevo_documento(self):
        Formulario(self)


if __name__ == "__main__":
    app = App()
    app.mainloop()